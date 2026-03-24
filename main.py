"""
404 Scanner - No API Key Version  (Phishing-Accurate Edition)
--------------------------------------------------------------
Fixes applied:
  1. URLhaus is a MALWARE db — added dedicated phishing checks via
     PhishTank (free, no key for read-only lookup) + OpenPhish feed.
  2. Rich heuristic engine covers typosquatting, subdomain abuse,
     homograph chars, high-entropy domains, redirect chains, etc.
  3. assess_risk() recalibrated — single strong indicator now raises
     MEDIUM; 2 → HIGH; 3+ → CRITICAL.
  4. URL scanner now also follows redirects and inspects final URL.
  5. Phishing keyword lists greatly expanded.
  6. Domain age / newly-registered check via WHOIS (python-whois).
  7. TLD reputation scoring added.
  8. All scan types updated to use the improved helpers.

Run:
    pip install requests python-whois Pillow python-docx PyPDF2 scapy
    python 404_scanner_fixed.py
"""

import hashlib
import ipaddress
import json
import math
import os
import re
import socket
import sys
import time
from datetime import datetime, timezone
from urllib.parse import urlparse

import requests

# ──────────────────────────────────────────────
# Styling / UI Helpers
# ──────────────────────────────────────────────

class Color:
    RED     = "\033[91m"
    ORANGE  = "\033[33m"
    YELLOW  = "\033[93m"
    GREEN   = "\033[92m"
    CYAN    = "\033[96m"
    BLUE    = "\033[94m"
    MAGENTA = "\033[95m"
    BOLD    = "\033[1m"
    RESET   = "\033[0m"

def banner():
    print(f"""{Color.CYAN}{Color.BOLD}

   ██╗  ██╗ ██████╗ ██╗  ██╗
   ██║  ██║██╔═████╗██║  ██║
   ███████║██║██╔██║███████║
   ╚════██║████╔╝██║╚════██║
        ██║╚██████╔╝     ██║
        ╚═╝ ╚═════╝      ╚═╝

        ███████╗ ██████╗ █████╗ ███╗   ██╗███╗   ██╗███████╗██████╗ 
        ██╔════╝██╔════╝██╔══██╗████╗  ██║████╗  ██║██╔════╝██╔══██╗
        ███████╗██║     ███████║██╔██╗ ██║██╔██╗ ██║█████╗  ██████╔╝
        ╚════██║██║     ██╔══██║██║╚██╗██║██║╚██╗██║██╔══╝  ██╔══██╗
        ███████║╚██████╗██║  ██║██║ ╚████║██║ ╚████║███████╗██║  ██║
        ╚══════╝ ╚═════╝╚═╝  ╚═╝╚═╝  ╚═══╝╚═╝  ╚═══╝╚══════╝╚═╝  ╚═╝

        404 Scanner  | Powered by Team 404  | Blue Team🛡️{Color.RESET}""")

def divider(char="─", length=65, color=Color.CYAN):
    print(f"{color}{char * length}{Color.RESET}")

def log_info(msg):    print(f"{Color.CYAN}[*]{Color.RESET} {msg}")
def log_success(msg): print(f"{Color.GREEN}[+]{Color.RESET} {msg}")
def log_warn(msg):    print(f"{Color.YELLOW}[!]{Color.RESET} {msg}")
def log_error(msg):   print(f"{Color.RED}[x]{Color.RESET} {msg}")


# ──────────────────────────────────────────────
# Helper Utilities
# ──────────────────────────────────────────────

def compute_sha256(filepath: str) -> str:
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def compute_md5(filepath: str) -> str:
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def is_private_ip(ip: str) -> bool:
    try:
        return ipaddress.ip_address(ip).is_private
    except ValueError:
        return False

def shannon_entropy(s: str) -> float:
    """Calculate Shannon entropy of a string (high = random/suspicious)."""
    if not s:
        return 0.0
    freq = {}
    for c in s:
        freq[c] = freq.get(c, 0) + 1
    length = len(s)
    return -sum((v / length) * math.log2(v / length) for v in freq.values())

def strip_www(domain: str) -> str:
    return re.sub(r"^www\.", "", domain.lower())


# ──────────────────────────────────────────────
# Risk Assessment  (recalibrated)
# ──────────────────────────────────────────────

# Each flag carries a weight.  Weights are summed; thresholds decide level.
# Default weight = 1.  High-confidence indicators have weight > 1.

HIGH_WEIGHT_INDICATORS = {
    # These phrases in a flag string bump its weight to 2
    "phishtank"             : 3,
    "openphish"             : 3,
    "urlhaus"               : 2,
    "malwarebazaar"         : 3,
    "currently active"      : 2,
    "active malicious"      : 2,
    "serving malware"       : 2,
    "malicious url"         : 2,
    "direct ip address url" : 2,
    "fake extension"        : 2,
    "javascript"            : 1,
    "openaction"            : 2,
    "launch action"         : 2,
}

def flag_weight(flag: str) -> int:
    fl = flag.lower()
    for keyword, weight in HIGH_WEIGHT_INDICATORS.items():
        if keyword in fl:
            return weight
    return 1

def assess_risk(flags: list) -> str:
    """
    Weighted scoring:
      0        → CLEAN
      1–2      → LOW
      3–4      → MEDIUM
      5–7      → HIGH
      8+       → CRITICAL
    """
    score = sum(flag_weight(f) for f in flags)
    if score == 0:   return "CLEAN"
    if score <= 2:   return "LOW"
    if score <= 4:   return "MEDIUM"
    if score <= 7:   return "HIGH"
    return "CRITICAL"

def risk_badge(level: str) -> str:
    badges = {
        "CRITICAL": f"{Color.RED}{Color.BOLD}[CRITICAL]{Color.RESET}",
        "HIGH":     f"{Color.ORANGE}{Color.BOLD}[HIGH]{Color.RESET}",
        "MEDIUM":   f"{Color.YELLOW}{Color.BOLD}[MEDIUM]{Color.RESET}",
        "LOW":      f"{Color.YELLOW}[LOW]{Color.RESET}",
        "CLEAN":    f"{Color.GREEN}{Color.BOLD}[CLEAN]{Color.RESET}",
    }
    return badges.get(level, f"{Color.GREEN}[CLEAN]{Color.RESET}")


# ──────────────────────────────────────────────
# Phishing / Threat Intelligence Helpers
# ──────────────────────────────────────────────

# ── OpenPhish feed (free, no key) ──────────────
_openphish_cache: set = set()
_openphish_loaded: bool = False

def _load_openphish():
    global _openphish_cache, _openphish_loaded
    if _openphish_loaded:
        return
    try:
        r = requests.get("https://openphish.com/feed.txt", timeout=12)
        if r.status_code == 200:
            _openphish_cache = {line.strip().lower() for line in r.text.splitlines() if line.strip()}
            log_info(f"OpenPhish feed loaded: {len(_openphish_cache)} entries")
        _openphish_loaded = True
    except Exception:
        _openphish_loaded = True   # don't retry on failure

def openphish_check(url: str) -> bool:
    """Return True if url (or its normalized form) is in the OpenPhish feed."""
    _load_openphish()
    url_lc = url.lower().rstrip("/")
    # exact match
    if url_lc in _openphish_cache:
        return True
    # domain-level match
    try:
        host = urlparse(url_lc).hostname or ""
        return any(host in entry for entry in _openphish_cache)
    except Exception:
        return False


# ── PhishTank (free, no key for basic lookup) ──
def phishtank_check(url: str) -> dict:
    """
    Check url against PhishTank's free API (no key = rate-limited but works).
    Returns dict: found (bool), verified (bool), online (bool).
    """
    result = {"found": False, "verified": False, "online": False}
    try:
        r = requests.post(
            "https://checkurl.phishtank.com/checkurl/",
            data={"url": url, "format": "json"},
            headers={"User-Agent": "404Scanner/2.0"},
            timeout=12,
        )
        data = r.json()
        results = data.get("results", {})
        if results.get("in_database"):
            result["found"]    = True
            result["verified"] = results.get("verified", False)
            result["online"]   = results.get("online", False)
    except Exception:
        pass
    return result


# ── URLhaus helpers ────────────────────────────

def urlhaus_check_url(url: str) -> dict:
    result = {"found": False, "active": False, "threat": "N/A", "tags": []}
    try:
        r = requests.post(
            "https://urlhaus-api.abuse.ch/v1/url/",
            data={"url": url}, timeout=10,
        )
        data = r.json()
        if data.get("query_status") == "ok":
            result["found"]  = True
            result["active"] = data.get("url_status", "") == "online"
            result["threat"] = data.get("threat", "N/A")
            result["tags"]   = data.get("tags") or []
    except Exception:
        pass
    return result


def urlhaus_check_host(host: str) -> dict:
    result = {"found": False, "online_count": 0, "total_count": 0, "tags": []}
    try:
        r = requests.post(
            "https://urlhaus-api.abuse.ch/v1/host/",
            data={"host": host}, timeout=10,
        )
        data = r.json()
        if data.get("query_status") == "is_host":
            urls   = data.get("urls", [])
            online = [u for u in urls if u.get("url_status") == "online"]
            tags   = set()
            for u in urls:
                tags.update(u.get("tags") or [])
            result["found"]        = True
            result["online_count"] = len(online)
            result["total_count"]  = len(urls)
            result["tags"]         = list(tags)
    except Exception:
        pass
    return result


# ──────────────────────────────────────────────
# Phishing Heuristic Engine
# ──────────────────────────────────────────────

# Expanded brand list for typosquatting / impersonation detection
BRAND_KEYWORDS = [
    "paypal", "amazon", "apple", "microsoft", "google", "facebook", "instagram",
    "twitter", "netflix", "spotify", "ebay", "chase", "wellsfargo", "bankofamerica",
    "citibank", "hsbc", "barclays", "lloyds", "natwest", "santander",
    "linkedin", "dropbox", "icloud", "outlook", "office365", "onedrive",
    "coinbase", "binance", "crypto", "blockchain", "wallet",
    "dhl", "fedex", "ups", "usps", "royalmail",
    "irs", "hmrc", "gov", "sbi", "hdfc", "icici", "axis",
    "whatsapp", "telegram", "signal",
    "steam", "roblox", "epicgames",
]

PHISHING_ACTION_WORDS = [
    "login", "log-in", "signin", "sign-in", "verify", "verification",
    "secure", "security", "confirm", "account", "update", "validate",
    "authenticate", "password", "reset", "recovery", "unlock",
    "suspend", "suspended", "limited", "unusual", "alert",
    "billing", "payment", "invoice", "refund", "reward",
    "click", "urgent", "immediately", "action-required", "action_required",
]

# TLD reputation: 0 = clean, higher = more abused
TLD_RISK = {
    ".tk": 3, ".ml": 3, ".ga": 3, ".cf": 3, ".gq": 3,
    ".xyz": 2, ".top": 2, ".click": 2, ".zip": 2, ".mov": 2,
    ".work": 2, ".live": 1, ".online": 1, ".site": 1,
    ".club": 1, ".info": 1, ".biz": 1, ".pw": 2,
}

# Unicode homograph characters mapped to their lookalikes
HOMOGRAPH_MAP = {
    "а": "a", "е": "e", "о": "o", "р": "p", "с": "c",
    "х": "x", "у": "y", "ј": "j", "ԁ": "d", "ⅼ": "l",
}

def _normalize_homograph(s: str) -> str:
    return "".join(HOMOGRAPH_MAP.get(c, c) for c in s)

def phishing_heuristics(url: str, final_url: str = None) -> list:
    """
    Returns a list of phishing flag strings for the given URL.
    final_url is the URL after following redirects (if available).
    """
    flags = []
    check_url = final_url or url
    check_url_lc = check_url.lower()

    try:
        parsed     = urlparse(check_url if check_url.startswith("http") else "http://" + check_url)
        orig_parsed = urlparse(url if url.startswith("http") else "http://" + url)
    except Exception:
        return flags

    host        = (parsed.hostname or "").lower()
    orig_host   = (orig_parsed.hostname or "").lower()
    full_url_lc = check_url_lc
    path        = (parsed.path or "").lower()
    query       = (parsed.query or "").lower()

    # ── 1. Direct IP address ──────────────────
    try:
        ipaddress.ip_address(host)
        flags.append("Direct IP address URL — no legitimate domain (suspicious)")
    except ValueError:
        pass

    # ── 2. Redirect to different domain ───────
    if final_url and orig_host and host and orig_host != host:
        flags.append(f"Redirect to different domain: {orig_host} → {host}")

    # ── 3. @ symbol trick ─────────────────────
    if "@" in check_url:
        flags.append("URL contains @ symbol — browser ignores everything before it (phishing trick)")

    # ── 4. Suspicious TLD ─────────────────────
    for tld, weight in TLD_RISK.items():
        if host.endswith(tld):
            severity = "High-risk" if weight >= 2 else "Suspicious"
            flags.append(f"{severity} TLD detected: {tld} (commonly used in phishing)")
            break

    # ── 5. Brand name in subdomain / path (not in apex domain) ──
    parts   = host.split(".")
    apex    = ".".join(parts[-2:]) if len(parts) >= 2 else host
    subdomain = ".".join(parts[:-2]) if len(parts) > 2 else ""
    for brand in BRAND_KEYWORDS:
        if brand in subdomain:
            flags.append(f"Brand name '{brand}' used in subdomain — likely impersonation")
            break
        if brand in path:
            flags.append(f"Brand name '{brand}' used in URL path — possible phishing")
            break
        # Brand in apex but apex is NOT the brand itself (e.g. paypal-verify.com)
        if brand in apex and not apex.startswith(brand + "."):
            flags.append(f"Brand name '{brand}' embedded in domain with extra text — typosquat/impersonation")
            break

    # ── 6. Phishing action words in domain/path ──
    action_hits = []
    for word in PHISHING_ACTION_WORDS:
        if word in host or word in path or word in query:
            action_hits.append(word)
    if len(action_hits) >= 2:
        flags.append(f"Multiple phishing action words in URL: {', '.join(action_hits[:5])}")
    elif len(action_hits) == 1:
        flags.append(f"Phishing action word in URL: '{action_hits[0]}'")

    # ── 7. URL shorteners ─────────────────────
    shorteners = [
        "bit.ly", "tinyurl.com", "t.co", "goo.gl", "ow.ly",
        "rb.gy", "cutt.ly", "is.gd", "buff.ly", "short.io",
        "tiny.cc", "bl.ink", "shorturl.at",
    ]
    for s in shorteners:
        if s in host:
            flags.append(f"URL shortener detected ({s}) — hides true destination")
            break

    # ── 8. Executable file extension ──────────
    bad_exts = [".exe", ".bat", ".ps1", ".vbs", ".js", ".jar",
                ".scr", ".com", ".pif", ".cmd", ".msi", ".dll", ".hta"]
    for ext in bad_exts:
        if path.endswith(ext) or query.endswith(ext):
            flags.append(f"URL points to executable file type: {ext}")
            break

    # ── 9. Unusually long domain ──────────────
    if len(host) > 50:
        flags.append(f"Unusually long domain ({len(host)} chars) — often used in phishing")

    # ── 10. Excessive subdomains ──────────────
    if host.count(".") >= 4:
        flags.append(f"Excessive subdomain depth ({host.count('.')} dots) — evasion technique")

    # ── 11. Homograph / Unicode attack ────────
    norm = _normalize_homograph(host)
    if norm != host:
        flags.append(f"Homograph/Unicode character attack detected in domain")

    # ── 12. High entropy domain ───────────────
    domain_only = apex.split(".")[0]
    ent = shannon_entropy(domain_only)
    if ent > 3.8 and len(domain_only) > 8:
        flags.append(f"High-entropy domain name (entropy={ent:.2f}) — looks randomly generated (DGA/phishing)")

    # ── 13. Mixed HTTPS/HTTP content / no HTTPS ──
    if full_url_lc.startswith("http://") and any(
        kw in full_url_lc for kw in ["login", "signin", "account", "secure", "verify", "banking"]
    ):
        flags.append("Sensitive page served over HTTP (not HTTPS) — credential theft risk")

    # ── 14. Data URI ──────────────────────────
    if check_url_lc.startswith("data:"):
        flags.append("Data URI detected — can embed malicious HTML/JS without a server")

    # ── 15. Double slashes / path confusion ───
    if re.search(r"https?://[^/]+//", check_url_lc):
        flags.append("Double-slash in URL path — evasion / confusion technique")

    # ── 16. Numeric-heavy domain ─────────────
    digits_in_domain = sum(c.isdigit() for c in domain_only)
    if len(domain_only) > 0 and digits_in_domain / len(domain_only) > 0.5:
        flags.append("Domain is mostly numeric — unusual for legitimate sites")

    return flags


# ──────────────────────────────────────────────
# Domain Age Check (python-whois)
# ──────────────────────────────────────────────

def check_domain_age(domain: str) -> list:
    """Return age-related flags. Empty list if whois not available or failed."""
    flags = []
    try:
        import whois
        w = whois.whois(domain)
        creation = w.creation_date
        if isinstance(creation, list):
            creation = creation[0]
        if creation:
            if hasattr(creation, "tzinfo") and creation.tzinfo is None:
                creation = creation.replace(tzinfo=timezone.utc)
            age_days = (datetime.now(timezone.utc) - creation).days
            if age_days < 30:
                flags.append(f"Domain registered VERY RECENTLY ({age_days} days ago) — high phishing risk")
            elif age_days < 180:
                flags.append(f"Domain registered recently ({age_days} days ago) — moderate risk")
    except ImportError:
        pass   # python-whois not installed — skip silently
    except Exception:
        pass
    return flags


# ──────────────────────────────────────────────
# Follow Redirects
# ──────────────────────────────────────────────

def follow_redirects(url: str, timeout: int = 10) -> str:
    """Follow up to 10 redirects and return the final URL."""
    try:
        r = requests.get(url, allow_redirects=True, timeout=timeout,
                         headers={"User-Agent": "Mozilla/5.0 (404Scanner/2.0)"},
                         stream=True)
        r.close()
        return r.url
    except Exception:
        return url


# ──────────────────────────────────────────────
# 1. IP Scanner
# ──────────────────────────────────────────────

def scan_ip(ip: str) -> dict:
    log_info(f"Looking up IP: {ip}")
    result = {"target": ip, "type": "IP", "flags": [], "info": {}}

    try:
        r = requests.get(f"https://ipwho.is/{ip}", timeout=10)
        r.raise_for_status()
        data = r.json()
        result["info"] = {
            "Country"  : data.get("country", "N/A"),
            "City"     : data.get("city", "N/A"),
            "Region"   : data.get("region", "N/A"),
            "ISP"      : data.get("connection", {}).get("isp", "N/A"),
            "ASN"      : data.get("connection", {}).get("asn", "N/A"),
            "Org"      : data.get("connection", {}).get("org", "N/A"),
            "Latitude" : data.get("latitude", "N/A"),
            "Longitude": data.get("longitude", "N/A"),
            "Timezone" : data.get("timezone", {}).get("id", "N/A"),
        }
        isp = str(data.get("connection", {}).get("isp", "")).lower()
        org = str(data.get("connection", {}).get("org", "")).lower()
        sus_kw = ["tor", "vpn", "proxy", "hosting", "datacenter",
                  "anonymous", "bulletproof", "ovh", "hetzner", "digitalocean"]
        for kw in sus_kw:
            if kw in isp or kw in org:
                result["flags"].append(f"Suspicious ISP/Org keyword: '{kw}'")
        log_success("IP info retrieved!")
    except requests.RequestException as e:
        log_warn(f"IPwho.is failed: {e}")

    host_data = urlhaus_check_host(ip)
    if host_data["found"]:
        if host_data["online_count"] > 0:
            result["flags"].append(
                f"URLhaus: {host_data['online_count']} ACTIVE malicious URL(s) hosted on this IP")
        else:
            result["flags"].append(
                f"URLhaus: {host_data['total_count']} historical malicious URL(s) on this IP")
        if host_data["tags"]:
            result["info"]["URLhaus Tags"] = ", ".join(host_data["tags"])
    else:
        result["info"]["URLhaus Host"] = "Not found in URLhaus database"

    result["risk"] = assess_risk(result["flags"])
    return result


# ──────────────────────────────────────────────
# 2. Domain Scanner
# ──────────────────────────────────────────────

def scan_domain(domain: str) -> dict:
    log_info(f"Scanning domain: {domain}")
    result = {"target": domain, "type": "Domain", "flags": [], "info": {}}
    clean_domain = strip_www(domain)

    # DNS
    try:
        ip = socket.gethostbyname(domain)
        result["info"]["Resolved IP"] = ip
        if is_private_ip(ip):
            result["flags"].append("Resolves to private/local IP address")
    except socket.gaierror:
        result["flags"].append("Domain does not resolve (possibly inactive or fake)")
        result["info"]["Resolved IP"] = "FAILED"

    # URLhaus
    host_data = urlhaus_check_host(domain)
    if host_data["found"]:
        if host_data["online_count"] > 0:
            result["flags"].append(
                f"URLhaus: {host_data['online_count']} ACTIVE malicious URL(s) on this domain")
        else:
            result["flags"].append(
                f"URLhaus: {host_data['total_count']} historical malicious URL(s) on this domain")
        if host_data["tags"]:
            result["info"]["Threat Tags"] = ", ".join(host_data["tags"])
        log_success("Domain found in URLhaus — threat detected!")
    else:
        result["info"]["URLhaus Host"] = "Not found in URLhaus database"

    # OpenPhish
    if openphish_check("http://" + domain):
        result["flags"].append("OpenPhish: domain found in active phishing feed")
        log_success("Domain found in OpenPhish feed!")

    # PhishTank
    pt = phishtank_check("http://" + domain)
    if pt["found"]:
        result["flags"].append(
            "PhishTank: domain confirmed as phishing" +
            (" (currently online)" if pt["online"] else " (offline/historical)")
        )

    # Heuristics on domain alone
    result["flags"].extend(phishing_heuristics("http://" + domain))

    # Domain age
    result["flags"].extend(check_domain_age(clean_domain))

    result["risk"] = assess_risk(result["flags"])
    return result


# ──────────────────────────────────────────────
# 3. URL Scanner  (now includes phishing checks)
# ──────────────────────────────────────────────

def scan_url(url: str) -> dict:
    log_info(f"Scanning URL: {url}")
    result = {"target": url, "type": "URL", "flags": [], "info": {}}

    lookup_url = url if url.startswith("http") else "http://" + url

    # ── Follow redirects ──────────────────────
    log_info("Following redirects...")
    final_url = follow_redirects(lookup_url)
    if final_url.lower() != lookup_url.lower():
        result["info"]["Final URL (after redirect)"] = final_url
        log_info(f"Redirected to: {final_url}")

    # ── URLhaus URL-level ──────────────────────
    url_data = urlhaus_check_url(lookup_url)
    if url_data["found"]:
        result["flags"].append("URL found in URLhaus malware database")
        result["info"]["URL Status"] = "online" if url_data["active"] else "offline/unknown"
        result["info"]["Threat"]     = url_data["threat"]
        if url_data["tags"]:
            result["info"]["Tags"] = ", ".join(url_data["tags"])
        if url_data["active"]:
            result["flags"].append("URL is currently ACTIVE and serving malware!")
        log_success("Malicious URL detected via URLhaus!")
    else:
        result["info"]["URLhaus URL"] = "Not found in URLhaus URL database"

    # ── URLhaus host-level ─────────────────────
    try:
        parsed = urlparse(lookup_url)
        host   = parsed.hostname or lookup_url.split("/")[0]
        if host:
            host_data = urlhaus_check_host(host)
            if host_data["found"]:
                if host_data["online_count"] > 0:
                    result["flags"].append(
                        f"Host has {host_data['online_count']} ACTIVE malicious URL(s) on URLhaus")
                else:
                    result["flags"].append(
                        f"Host has {host_data['total_count']} historical malicious URL(s) on URLhaus")
                if host_data["tags"]:
                    result["info"]["Host Threat Tags"] = ", ".join(host_data["tags"])
            else:
                result["info"]["URLhaus Host"] = "Not found in host database"
    except Exception as e:
        log_warn(f"URLhaus host check failed: {e}")

    # ── OpenPhish ─────────────────────────────
    log_info("Checking OpenPhish phishing feed...")
    if openphish_check(lookup_url):
        result["flags"].append("OpenPhish: URL/domain found in active phishing feed!")
        log_success("Phishing URL confirmed via OpenPhish!")
    else:
        result["info"]["OpenPhish"] = "Not found in OpenPhish feed"

    # ── PhishTank ─────────────────────────────
    log_info("Checking PhishTank database...")
    pt = phishtank_check(lookup_url)
    if pt["found"]:
        result["flags"].append(
            "PhishTank: URL confirmed as phishing site" +
            (" (currently active)" if pt["online"] else " (offline/historical)")
        )
        log_success("Phishing URL confirmed via PhishTank!")
    else:
        result["info"]["PhishTank"] = "Not found in PhishTank database"

    # ── Phishing heuristics ───────────────────
    log_info("Running phishing heuristic analysis...")
    heuristic_flags = phishing_heuristics(lookup_url, final_url if final_url != lookup_url else None)
    result["flags"].extend(heuristic_flags)
    if heuristic_flags:
        log_warn(f"Heuristic analysis found {len(heuristic_flags)} indicator(s)")

    # ── Domain age ────────────────────────────
    try:
        parsed = urlparse(lookup_url)
        host = parsed.hostname or ""
        if host and not re.match(r"^\d{1,3}(\.\d{1,3}){3}$", host):
            result["flags"].extend(check_domain_age(strip_www(host)))
    except Exception:
        pass

    result["risk"] = assess_risk(result["flags"])

    # ── Summary line ─────────────────────────
    if result["risk"] == "CLEAN":
        result["info"]["Assessment"] = "No indicators found — appears clean (not guaranteed safe)"
    elif result["risk"] == "LOW":
        result["info"]["Assessment"] = "Minor indicators found — treat with caution"
    elif result["risk"] == "MEDIUM":
        result["info"]["Assessment"] = "Multiple phishing indicators — likely suspicious"
    elif result["risk"] in ("HIGH", "CRITICAL"):
        result["info"]["Assessment"] = "Strong evidence of phishing/malware — DO NOT visit"

    return result


# ──────────────────────────────────────────────
# 4. File Hash Scanner
# ──────────────────────────────────────────────

def scan_hash(file_hash: str) -> dict:
    log_info(f"Looking up hash: {file_hash}")
    result = {"target": file_hash, "type": "Hash", "flags": [], "info": {}}
    try:
        r = requests.post(
            "https://mb-api.abuse.ch/api/v1/",
            data={"query": "get_info", "hash": file_hash}, timeout=15,
        )
        data = r.json()
        status = data.get("query_status", "")
        if status == "ok":
            details = data.get("data", [{}])[0]
            result["flags"].append("Hash found in MalwareBazaar database")
            result["info"]["File Name"]  = details.get("file_name", "N/A")
            result["info"]["File Type"]  = details.get("file_type", "N/A")
            result["info"]["File Size"]  = f"{details.get('file_size', 'N/A')} bytes"
            result["info"]["Signature"]  = details.get("signature", "N/A")
            result["info"]["First Seen"] = details.get("first_seen", "N/A")
            result["info"]["Last Seen"]  = details.get("last_seen", "N/A")
            tags = details.get("tags") or []
            if tags:
                result["info"]["Tags"] = ", ".join(tags)
            result["info"]["Reporter"]   = details.get("reporter", "N/A")
            log_success("Hash found — MALWARE detected!")
        elif status == "hash_not_found":
            result["info"]["MalwareBazaar"] = "Hash not found (not known malware)"
            log_success("Hash lookup done — not found in MalwareBazaar.")
    except Exception as e:
        log_warn(f"MalwareBazaar lookup failed: {e}")
    result["risk"] = assess_risk(result["flags"])
    return result


# ──────────────────────────────────────────────
# 5. File Scanner
# ──────────────────────────────────────────────

def scan_file(filepath: str) -> dict:
    if not os.path.isfile(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    sha256 = compute_sha256(filepath)
    md5    = compute_md5(filepath)
    size   = os.path.getsize(filepath)
    log_info(f"SHA-256 : {sha256}")
    log_info(f"MD5     : {md5}")
    log_info(f"Size    : {size} bytes")
    result = scan_hash(sha256)
    result["target"] = filepath
    result["type"]   = "File"
    result["info"]["SHA-256"]   = sha256
    result["info"]["MD5"]       = md5
    result["info"]["File Size"] = f"{size} bytes"
    suspicious_exts = [".exe", ".bat", ".ps1", ".vbs", ".js", ".jar",
                       ".scr", ".com", ".pif", ".cmd", ".msi", ".dll"]
    ext = os.path.splitext(filepath)[1].lower()
    if ext in suspicious_exts:
        result["flags"].append(f"Suspicious file extension: {ext}")
    result["risk"] = assess_risk(result["flags"])
    return result


# ──────────────────────────────────────────────
# 6. Port Scanner
# ──────────────────────────────────────────────

COMMON_PORTS = {
    21: "FTP", 22: "SSH", 23: "Telnet", 25: "SMTP",
    53: "DNS", 80: "HTTP", 110: "POP3", 143: "IMAP",
    443: "HTTPS", 445: "SMB", 3306: "MySQL", 3389: "RDP",
    5900: "VNC", 6379: "Redis", 8080: "HTTP-Alt", 8443: "HTTPS-Alt",
    27017: "MongoDB",
}
RISKY_PORTS = {
    23: "Telnet (unencrypted)", 445: "SMB (ransomware target)",
    3389: "RDP (brute force risk)", 5900: "VNC (remote access)",
    6379: "Redis (often exposed)", 27017: "MongoDB (often exposed)",
}

def scan_ports(host: str) -> dict:
    log_info(f"Port scanning: {host}")
    result = {"target": host, "type": "Port Scan", "flags": [], "info": {}, "open_ports": []}
    try:
        ip = socket.gethostbyname(host)
        result["info"]["Resolved IP"] = ip
    except socket.gaierror:
        log_error(f"Cannot resolve host: {host}")
        result["risk"] = "UNKNOWN"
        return result
    open_ports = []
    log_info(f"Scanning {len(COMMON_PORTS)} common ports...")
    for port, service in COMMON_PORTS.items():
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(0.5)
            if sock.connect_ex((ip, port)) == 0:
                open_ports.append((port, service))
                if port in RISKY_PORTS:
                    result["flags"].append(
                        f"Risky port open: {port}/{service} — {RISKY_PORTS[port]}")
            sock.close()
        except Exception:
            pass
    result["open_ports"]         = open_ports
    result["info"]["Open Ports"] = len(open_ports)
    result["risk"]               = assess_risk(result["flags"])
    log_success(f"Found {len(open_ports)} open port(s).")
    return result


# ──────────────────────────────────────────────
# 7. PCAP Analyzer
# ──────────────────────────────────────────────

def scan_pcap(filepath: str) -> dict:
    try:
        from scapy.all import rdpcap, IP, IPv6, DNS, DNSQR, Raw
    except ImportError:
        raise ImportError("scapy not installed. Run: pip install scapy")
    if not os.path.isfile(filepath):
        raise FileNotFoundError(f"PCAP file not found: {filepath}")
    log_info(f"Reading PCAP: {filepath}")
    packets = rdpcap(filepath)
    log_success(f"Loaded {len(packets)} packets.")
    ips     = set()
    domains = set()
    skip    = ("10.", "192.168.", "127.", "0.", "255.", "::1", "fe80")
    for pkt in packets:
        for layer in (IP, IPv6):
            if pkt.haslayer(layer):
                for addr in (str(pkt[layer].src), str(pkt[layer].dst)):
                    if not any(addr.startswith(p) for p in skip):
                        ips.add(addr)
        if pkt.haslayer(DNS) and pkt.haslayer(DNSQR):
            try:
                qname = pkt[DNSQR].qname.decode("utf-8", errors="ignore").rstrip(".")
                if qname and "." in qname and not qname.endswith(".local"):
                    domains.add(qname)
            except Exception:
                pass
        if pkt.haslayer(Raw):
            try:
                payload = pkt[Raw].load.decode("utf-8", errors="ignore")
                for match in re.findall(r"Host:\s*([^\r\n]+)", payload):
                    host = match.strip().split(":")[0]
                    if "." in host:
                        domains.add(host)
            except Exception:
                pass
    log_info(f"Extracted {len(ips)} IPs and {len(domains)} domains.")
    ip_results     = []
    domain_results = []
    print(f"\n{Color.BOLD}  [ Scanning IPs ]{Color.RESET}")
    divider()
    for ip in sorted(ips):
        r      = scan_ip(ip)
        badge  = risk_badge(r["risk"])
        country = r["info"].get("Country", "N/A")
        isp     = r["info"].get("ISP", "N/A")
        print(f"  {Color.CYAN}{ip:20}{Color.RESET} {badge} | {country} | {isp}")
        for flag in r["flags"]:
            print(f"      {Color.YELLOW}>> {flag}{Color.RESET}")
        ip_results.append(r)
        time.sleep(0.3)
    print(f"\n{Color.BOLD}  [ Scanning Domains ]{Color.RESET}")
    divider()
    for domain in sorted(domains):
        r     = scan_domain(domain)
        badge = risk_badge(r["risk"])
        print(f"  {Color.CYAN}{domain:40}{Color.RESET} {badge}")
        for flag in r["flags"]:
            print(f"      {Color.YELLOW}>> {flag}{Color.RESET}")
        domain_results.append(r)
        time.sleep(0.3)
    threats = sum(1 for r in ip_results + domain_results if r["risk"] != "CLEAN")
    return {
        "target": filepath, "type": "PCAP",
        "ip_results": ip_results, "domain_results": domain_results,
        "threats": threats, "flags": [],
        "info": {"IPs Scanned": len(ips), "Domains Scanned": len(domains)},
        "risk": (
            "CRITICAL" if threats >= 5 else
            "HIGH"     if threats >= 2 else
            "MEDIUM"   if threats >= 1 else "CLEAN"
        ),
    }


# ──────────────────────────────────────────────
# 8. PDF Scanner
# ──────────────────────────────────────────────

def scan_pdf(filepath: str) -> dict:
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    if not os.path.isfile(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    result = {
        "target": filepath, "type": "PDF",
        "flags": [], "info": {}, "urls": [], "url_results": [],
    }
    sha256 = compute_sha256(filepath)
    md5    = compute_md5(filepath)
    result["info"]["SHA-256"]   = sha256
    result["info"]["MD5"]       = md5
    result["info"]["File Size"] = f"{os.path.getsize(filepath)} bytes"
    log_info("Checking PDF hash on MalwareBazaar...")
    try:
        r = requests.post("https://mb-api.abuse.ch/api/v1/",
                          data={"query": "get_info", "hash": sha256}, timeout=15)
        data = r.json()
        if data.get("query_status") == "ok":
            result["flags"].append("PDF hash found in MalwareBazaar malware database!")
            details = data.get("data", [{}])[0]
            result["info"]["MalwareBazaar"] = details.get("signature", "Known malware")
    except Exception:
        pass
    log_info("Parsing PDF structure...")
    try:
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
        result["info"]["Pages"]     = len(reader.pages)
        result["info"]["Encrypted"] = "Yes" if reader.is_encrypted else "No"
        meta = reader.metadata
        if meta:
            result["info"]["Author"]   = meta.get("/Author",   "N/A")
            result["info"]["Creator"]  = meta.get("/Creator",  "N/A")
            result["info"]["Producer"] = meta.get("/Producer", "N/A")
            result["info"]["Created"]  = meta.get("/CreationDate", "N/A")
        full_text = ""
        for page in reader.pages:
            try:
                full_text += page.extract_text() or ""
            except Exception:
                pass
        urls_found = re.findall(r'https?://[^\s\]\[<>"\'{}|\\^`)(,;]+', full_text)
        result["urls"] = list(set(urls_found))
        if result["urls"]:
            result["info"]["URLs Found"] = len(result["urls"])
            result["flags"].append(f"{len(result['urls'])} URL(s) embedded in PDF")
        phishing_words = [
            "verify your account", "click here", "login", "confirm your",
            "update your", "suspended", "urgent", "password", "credit card",
            "social security", "wire transfer", "enable macros",
            "your account has been", "action required",
        ]
        text_lower = full_text.lower()
        found_kw = [w for w in phishing_words if w in text_lower]
        if found_kw:
            result["flags"].append(f"Phishing keywords found: {', '.join(found_kw)}")
        log_success("PDF structure parsed!")
    except Exception as e:
        log_warn(f"PDF parsing error: {e}")
        result["flags"].append("PDF is corrupted or malformed (possible evasion technique)")
    log_info("Scanning PDF raw content for dangerous objects...")
    try:
        with open(filepath, "rb") as f:
            raw = f.read().decode("latin-1", errors="ignore")
        dangerous_keys = {
            "/JavaScript"  : "JavaScript code detected inside PDF",
            "/JS"          : "JavaScript (short form) detected inside PDF",
            "/AA"          : "Auto Action detected (runs on open)",
            "/OpenAction"  : "OpenAction detected (auto-executes on open)",
            "/Launch"      : "Launch action detected (can run executables)",
            "/EmbeddedFile": "Embedded file detected inside PDF",
            "/RichMedia"   : "RichMedia object detected (Flash/video embed)",
            "/XFA"         : "XFA form detected (commonly used in exploits)",
            "/Encrypt"     : "Encryption object found",
            "/AcroForm"    : "AcroForm detected (interactive form)",
        }
        found_keys = []
        for key, msg in dangerous_keys.items():
            if key in raw:
                found_keys.append(key)
                result["flags"].append(msg)
        result["info"]["Dangerous Objects"] = ", ".join(found_keys) if found_keys else "None"
    except Exception as e:
        log_warn(f"Raw byte scan error: {e}")
    if result["urls"]:
        log_info(f"Scanning {len(result['urls'])} embedded URL(s) for phishing/malware...")
        for url in result["urls"][:10]:
            url_flags  = phishing_heuristics(url)
            url_data   = urlhaus_check_url(url)
            pt         = phishtank_check(url)
            op         = openphish_check(url)
            is_bad     = url_data["found"] or pt["found"] or op or bool(url_flags)
            status_label = "MALICIOUS" if (url_data["found"] or pt["found"] or op) else \
                           "SUSPICIOUS" if url_flags else "CLEAN"
            if is_bad and status_label in ("MALICIOUS", "SUSPICIOUS"):
                result["flags"].append(f"{status_label} URL in PDF: {url}")
            result["url_results"].append({"url": url, "status": status_label})
            time.sleep(0.2)
    result["risk"] = assess_risk(result["flags"])
    return result


# ──────────────────────────────────────────────
# 9. Image Scanner
# ──────────────────────────────────────────────

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".ico"}

def scan_image(filepath: str) -> dict:
    if not os.path.isfile(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    ext    = os.path.splitext(filepath)[1].lower()
    result = {"target": filepath, "type": "Image", "flags": [], "info": {}, "url_results": []}
    result["info"]["File Size"] = f"{os.path.getsize(filepath)} bytes"
    result["info"]["Extension"] = ext if ext else "Unknown"
    sha256 = compute_sha256(filepath)
    md5    = compute_md5(filepath)
    result["info"]["SHA-256"] = sha256
    result["info"]["MD5"]     = md5
    log_info("Checking image hash on MalwareBazaar...")
    try:
        r = requests.post("https://mb-api.abuse.ch/api/v1/",
                          data={"query": "get_info", "hash": sha256}, timeout=15)
        data = r.json()
        if data.get("query_status") == "ok":
            result["flags"].append("Image hash found in MalwareBazaar database!")
            details = data.get("data", [{}])[0]
            result["info"]["MalwareBazaar"] = details.get("signature", "Known malware")
    except Exception:
        pass
    log_info("Checking file magic bytes...")
    magic_signatures = {
        b"\xff\xd8\xff"     : "JPEG",
        b"\x89PNG\r\n\x1a\n": "PNG",
        b"GIF87a"           : "GIF",
        b"GIF89a"           : "GIF",
        b"BM"               : "BMP",
        b"RIFF"             : "WEBP",
    }
    try:
        with open(filepath, "rb") as f:
            header = f.read(12)
        detected = None
        for sig, fmt in magic_signatures.items():
            if header.startswith(sig):
                detected = fmt
                break
        result["info"]["Detected Format"] = detected or "Unknown/Suspicious"
        if detected is None:
            result["flags"].append("File magic bytes do not match any known image format")
        elif ext in IMAGE_EXTENSIONS:
            ext_map = {".jpg": "JPEG", ".jpeg": "JPEG", ".png": "PNG",
                       ".gif": "GIF", ".bmp": "BMP", ".webp": "WEBP"}
            expected = ext_map.get(ext)
            if expected and detected != expected:
                result["flags"].append(
                    f"Fake extension! File is {detected} but saved as {ext}")
    except Exception as e:
        log_warn(f"Magic byte check failed: {e}")
    log_info("Extracting EXIF metadata...")
    try:
        from PIL import Image as PILImage
        from PIL.ExifTags import TAGS
        img = PILImage.open(filepath)
        result["info"]["Image Size"] = f"{img.width} x {img.height} px"
        result["info"]["Mode"]       = img.mode
        exif_data = img._getexif() if hasattr(img, "_getexif") else None
        if exif_data:
            for tag_id, value in exif_data.items():
                tag = TAGS.get(tag_id, tag_id)
                if tag in ("Make", "Model", "Software", "DateTime",
                           "Artist", "Copyright", "GPSInfo"):
                    result["info"][f"EXIF {tag}"] = str(value)[:80]
            if "GPSInfo" in {TAGS.get(t) for t in exif_data}:
                result["flags"].append("GPS location data found in EXIF metadata!")
    except ImportError:
        log_warn("Pillow not installed — skipping EXIF. Run: pip install Pillow")
    except Exception:
        pass
    log_info("Scanning for hidden strings in image data...")
    try:
        with open(filepath, "rb") as f:
            raw = f.read().decode("latin-1", errors="ignore")
        urls_found = re.findall(r'https?://[^\s\x00-\x1f\x7f-\xff]{8,}', raw)
        if urls_found:
            result["flags"].append(f"{len(urls_found)} URL(s) found hidden in image data")
            for u in urls_found[:10]:
                hf = phishing_heuristics(u)
                status = "SUSPICIOUS" if hf else "UNKNOWN"
                result["url_results"].append({"url": u, "status": status})
        sus_strings = ["eval(", "base64", "powershell", "cmd.exe",
                       "wget ", "curl ", "/bin/sh", "chmod +x"]
        for s in sus_strings:
            if s.lower() in raw.lower():
                result["flags"].append(f"Suspicious string found in image: '{s}'")
    except Exception:
        pass
    result["risk"] = assess_risk(result["flags"])
    log_success("Image scan complete!")
    return result


# ──────────────────────────────────────────────
# 10. Word File Scanner
# ──────────────────────────────────────────────

def scan_word(filepath: str) -> dict:
    if not os.path.isfile(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    result = {"target": filepath, "type": "Word", "flags": [], "info": {}, "url_results": []}
    result["info"]["File Size"] = f"{os.path.getsize(filepath)} bytes"
    sha256 = compute_sha256(filepath)
    md5    = compute_md5(filepath)
    result["info"]["SHA-256"] = sha256
    result["info"]["MD5"]     = md5
    log_info("Checking Word file hash on MalwareBazaar...")
    try:
        r = requests.post("https://mb-api.abuse.ch/api/v1/",
                          data={"query": "get_info", "hash": sha256}, timeout=15)
        data = r.json()
        if data.get("query_status") == "ok":
            result["flags"].append("Word file hash found in MalwareBazaar database!")
            details = data.get("data", [{}])[0]
            result["info"]["MalwareBazaar"] = details.get("signature", "Known malware")
    except Exception:
        pass
    log_info("Parsing Word document structure...")
    try:
        import docx
        doc = docx.Document(filepath)
        props = doc.core_properties
        result["info"]["Author"]     = props.author   or "N/A"
        result["info"]["Created"]    = str(props.created)  if props.created  else "N/A"
        result["info"]["Modified"]   = str(props.modified) if props.modified else "N/A"
        result["info"]["Company"]    = props.company  or "N/A"
        result["info"]["Revision"]   = str(props.revision) if props.revision else "N/A"
        result["info"]["Paragraphs"] = str(len(doc.paragraphs))
        full_text  = "\n".join([p.text for p in doc.paragraphs])
        urls_found = re.findall(r'https?://[^\s\]\[<>"\']+', full_text)
        if urls_found:
            result["flags"].append(f"{len(urls_found)} URL(s) found in document")
            result["info"]["URLs Found"] = len(urls_found)
        phishing_words = [
            "verify your account", "click here", "login", "confirm your",
            "update your", "suspended", "urgent", "password",
            "wire transfer", "invoice", "enable macros",
            "enable content", "enable editing", "action required",
        ]
        text_lower = full_text.lower()
        found_kw = [w for w in phishing_words if w in text_lower]
        if found_kw:
            result["flags"].append(
                f"Phishing/social-engineering keywords: {', '.join(found_kw)}")
        for url in list(set(urls_found))[:10]:
            hf         = phishing_heuristics(url)
            url_data   = urlhaus_check_url(url)
            pt         = phishtank_check(url)
            op         = openphish_check(url)
            is_confirmed = url_data["found"] or pt["found"] or op
            status_label = "MALICIOUS"   if is_confirmed else \
                           "SUSPICIOUS"  if hf           else "CLEAN"
            if status_label in ("MALICIOUS", "SUSPICIOUS"):
                result["flags"].append(f"{status_label} URL in document: {url}")
            result["url_results"].append({"url": url, "status": status_label})
            time.sleep(0.2)
        log_success("Word document parsed!")
    except ImportError:
        log_warn("python-docx not installed. Run: pip install python-docx")
        result["flags"].append("Could not parse .docx — python-docx not installed")
    except Exception as e:
        log_warn(f"Word parse error: {e}")
    log_info("Scanning for macros and embedded objects...")
    try:
        import zipfile
        if zipfile.is_zipfile(filepath):
            with zipfile.ZipFile(filepath, "r") as z:
                names = z.namelist()
                result["info"]["Internal Files"] = len(names)
                macro_files = [n for n in names if "vba" in n.lower() or "macro" in n.lower()]
                if macro_files:
                    result["flags"].append(f"VBA Macro detected: {', '.join(macro_files)}")
                ole_files = [n for n in names if n.endswith((".bin", ".ole"))]
                if ole_files:
                    result["flags"].append(
                        f"Embedded OLE object(s) found: {len(ole_files)}")
                for name in names:
                    if "rels" in name:
                        content = z.read(name).decode("utf-8", errors="ignore")
                        ext_links = re.findall(r'Target="(https?://[^"]+)"', content)
                        if ext_links:
                            result["flags"].append(
                                f"External link in document relationships: {ext_links[0]}")
        else:
            result["flags"].append(
                "File is not a valid .docx (not a ZIP archive) — possibly .doc or corrupt")
    except Exception as e:
        log_warn(f"Raw scan error: {e}")
    result["risk"] = assess_risk(result["flags"])
    return result


# ──────────────────────────────────────────────
# 11. TXT File Scanner
# ──────────────────────────────────────────────

def scan_txt(filepath: str) -> dict:
    if not os.path.isfile(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    result = {
        "target": filepath, "type": "TXT",
        "flags": [], "info": {}, "url_results": [], "ip_results": [],
    }
    result["info"]["File Size"] = f"{os.path.getsize(filepath)} bytes"
    sha256 = compute_sha256(filepath)
    md5    = compute_md5(filepath)
    result["info"]["SHA-256"] = sha256
    result["info"]["MD5"]     = md5
    log_info("Checking TXT file hash on MalwareBazaar...")
    try:
        r = requests.post("https://mb-api.abuse.ch/api/v1/",
                          data={"query": "get_info", "hash": sha256}, timeout=15)
        data = r.json()
        if data.get("query_status") == "ok":
            result["flags"].append("File hash found in MalwareBazaar database!")
    except Exception:
        pass
    log_info("Reading and analyzing TXT content...")
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
    except Exception as e:
        log_error(f"Could not read file: {e}")
        result["risk"] = "UNKNOWN"
        return result
    result["info"]["Lines"]      = content.count("\n")
    result["info"]["Characters"] = len(content)
    urls_found = list(set(re.findall(r'https?://[^\s\]\[<>"\']+', content)))
    if urls_found:
        result["info"]["URLs Found"] = len(urls_found)
        result["flags"].append(f"{len(urls_found)} URL(s) found in file")
        log_info(f"Scanning {min(len(urls_found), 10)} URL(s)...")
        for url in urls_found[:10]:
            hf         = phishing_heuristics(url)
            url_data   = urlhaus_check_url(url)
            pt         = phishtank_check(url)
            op         = openphish_check(url)
            is_confirmed = url_data["found"] or pt["found"] or op
            status_label = "MALICIOUS"   if is_confirmed else \
                           "SUSPICIOUS"  if hf           else "CLEAN"
            if status_label in ("MALICIOUS", "SUSPICIOUS"):
                result["flags"].append(f"{status_label} URL found: {url}")
            result["url_results"].append({"url": url, "status": status_label})
            time.sleep(0.2)
    ips_found  = list(set(re.findall(
        r'\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b', content)))
    public_ips = [ip for ip in ips_found if not is_private_ip(ip)]
    if public_ips:
        result["info"]["IPs Found"] = len(public_ips)
        result["flags"].append(f"{len(public_ips)} public IP address(es) found in file")
        for ip in public_ips[:5]:
            ip_res = scan_ip(ip)
            result["ip_results"].append(ip_res)
            if ip_res["risk"] != "CLEAN":
                result["flags"].append(f"Suspicious IP in file: {ip} [{ip_res['risk']}]")
            time.sleep(0.3)
    script_patterns = [
        (r"powershell\s+-[eE]",        "Encoded PowerShell command detected"),
        (r"cmd\.exe\s*/[cC]",          "CMD execution command detected"),
        (r"base64[_\-]?decode",        "Base64 decode pattern detected"),
        (r"eval\s*\(",                  "eval() function detected"),
        (r"wget\s+http",               "wget download command detected"),
        (r"curl\s+http",               "curl download command detected"),
        (r"chmod\s+\+x",               "chmod +x (Linux execution) detected"),
        (r"/bin/bash|/bin/sh",         "Shell execution path detected"),
        (r"nc\s+-[lvp]",               "Netcat listener command detected"),
        (r"msfvenom|meterpreter|msf>", "Metasploit payload string detected"),
        (r"[A-Za-z0-9+/]{100,}={0,2}", "Long Base64 encoded string detected"),
    ]
    for pattern, msg in script_patterns:
        if re.search(pattern, content, re.IGNORECASE):
            result["flags"].append(msg)
    sus_keywords = ["password:", "passwd:", "secret:", "api_key:",
                    "private_key", "BEGIN RSA", "BEGIN CERTIFICATE",
                    "access_token", "authorization:"]
    content_lower = content.lower()
    for kw in sus_keywords:
        if kw.lower() in content_lower:
            result["flags"].append(f"Sensitive data keyword found: '{kw}'")
    result["risk"] = assess_risk(result["flags"])
    log_success("TXT scan complete!")
    return result


# ──────────────────────────────────────────────
# Report Printer (Terminal)
# ──────────────────────────────────────────────

def print_report(result: dict):
    print()
    divider("=")
    print(f"{Color.BOLD}  404 SCANNER - SCAN REPORT  |  Team 404{Color.RESET}")
    print(f"  Type   : {Color.CYAN}{result['type'].upper()}{Color.RESET}")
    print(f"  Target : {Color.MAGENTA}{result['target']}{Color.RESET}")
    divider("=")
    print(f"\n  Risk Level : {risk_badge(result['risk'])}")
    if result["flags"]:
        print(f"\n  {Color.BOLD}Threat Indicators ({len(result['flags'])}):{Color.RESET}")
        for flag in result["flags"]:
            w = flag_weight(flag)
            weight_label = f" {Color.RED}[w={w}]{Color.RESET}" if w > 1 else ""
            print(f"    {Color.YELLOW}>>{Color.RESET} {flag}{weight_label}")
    else:
        print(f"\n  {Color.GREEN}No threat indicators found.{Color.RESET}")
    if result.get("info"):
        print(f"\n  {Color.BOLD}Details:{Color.RESET}")
        for k, v in result["info"].items():
            print(f"    {k:28}: {v}")
    if result.get("open_ports"):
        print(f"\n  {Color.BOLD}Open Ports:{Color.RESET}")
        for port, service in result["open_ports"]:
            color = Color.RED if port in RISKY_PORTS else Color.GREEN
            print(f"    {color}{port:6}{Color.RESET}  {service}")
    if result.get("url_results"):
        print(f"\n  {Color.BOLD}Embedded URLs ({len(result['url_results'])}):{Color.RESET}")
        for item in result["url_results"]:
            color = (Color.RED    if item["status"] == "MALICIOUS"  else
                     Color.YELLOW if item["status"] == "SUSPICIOUS" else Color.GREEN)
            print(f"    {color}[{item['status']:9}]{Color.RESET}  {item['url']}")
    if result.get("ip_results"):
        print(f"\n  {Color.BOLD}IPs Found in File:{Color.RESET}")
        for r in result["ip_results"]:
            badge   = risk_badge(r["risk"])
            country = r["info"].get("Country", "N/A")
            print(f"    {Color.CYAN}{r['target']:18}{Color.RESET} {badge} | {country}")
    divider("=")
    print()


# ──────────────────────────────────────────────
# Report Generation (File)
# ──────────────────────────────────────────────

def generate_report(result: dict) -> str:
    now   = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    lines = []
    lines.append("=" * 65)
    lines.append("          404 SCANNER - SCAN REPORT")
    lines.append("                  Team 404")
    lines.append("       Phishing-Accurate Edition v2.0")
    lines.append("=" * 65)
    lines.append(f"  Date        : {now}")
    lines.append(f"  Scan Type   : {result['type'].upper()}")
    lines.append(f"  Target      : {result['target']}")
    lines.append(f"  Risk Level  : {result['risk']}")
    lines.append("")
    if result["flags"]:
        lines.append("-" * 65)
        lines.append(f"  THREAT INDICATORS ({len(result['flags'])})")
        lines.append("-" * 65)
        for flag in result["flags"]:
            w = flag_weight(flag)
            lines.append(f"  >> [w={w}] {flag}")
        lines.append("")
    if result.get("info"):
        lines.append("-" * 65)
        lines.append("  DETAILS")
        lines.append("-" * 65)
        for k, v in result["info"].items():
            lines.append(f"  {k:28}: {v}")
        lines.append("")
    if result.get("open_ports"):
        lines.append("-" * 65)
        lines.append("  OPEN PORTS")
        lines.append("-" * 65)
        for port, service in result["open_ports"]:
            risky = " [RISKY]" if port in RISKY_PORTS else ""
            lines.append(f"  {port:6}  {service}{risky}")
        lines.append("")
    if result.get("url_results"):
        lines.append("-" * 65)
        lines.append(f"  EMBEDDED URLS ({len(result['url_results'])})")
        lines.append("-" * 65)
        for item in result["url_results"]:
            lines.append(f"  [{item['status']:9}]  {item['url']}")
        lines.append("")
    if result.get("ip_results") or result.get("domain_results"):
        lines.append("-" * 65)
        lines.append("  PCAP RESULTS")
        lines.append("-" * 65)
        for r in result.get("ip_results", []):
            lines.append(f"  {r['target']:20}  Risk: {r['risk']:10}  "
                         f"Country:{r['info'].get('Country','N/A')}  ISP:{r['info'].get('ISP','N/A')}")
            for flag in r["flags"]:
                lines.append(f"      >> {flag}")
        for r in result.get("domain_results", []):
            lines.append(f"  {r['target']:40}  Risk: {r['risk']}")
            for flag in r["flags"]:
                lines.append(f"      >> {flag}")
        lines.append(f"  Total Threats Found : {result.get('threats', 0)}")
        lines.append("")
    lines.append("=" * 65)
    lines.append("  Generated by 404 Scanner | Team 404 | v2.0")
    lines.append("=" * 65)
    return "\n".join(lines)

def ask_save_report(result: dict):
    ans = input(f"\n{Color.CYAN}Save report to file? (y/n): {Color.RESET}").strip().lower()
    if ans == "y":
        content   = generate_report(result)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename  = f"404scan_report_{timestamp}.txt"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(content)
        log_success(f"Report saved: {Color.MAGENTA}{filename}{Color.RESET}")
    else:
        log_info("Report not saved.")


# ──────────────────────────────────────────────
# Interactive Menu
# ──────────────────────────────────────────────

MENU_OPTIONS = [
    ("[1]  URL       ", "Scan a URL (URLhaus + PhishTank + OpenPhish + heuristics)"),
    ("[2]  IP        ", "Scan an IP address (IPwho.is + URLhaus)"),
    ("[3]  Domain    ", "Scan a domain (DNS + URLhaus + PhishTank + OpenPhish)"),
    ("[4]  Hash      ", "Look up a file hash (MalwareBazaar)"),
    ("[5]  File      ", "Scan a local file (hash + extension check)"),
    ("[6]  PDF       ", "Scan a PDF file (JS, URLs, embedded objects)"),
    ("[7]  Image     ", "Scan an image file (EXIF, fake ext, hidden data)"),
    ("[8]  Word      ", "Scan a Word .docx file (macros, URLs, metadata)"),
    ("[9]  TXT       ", "Scan a TXT file (URLs, IPs, script patterns)"),
    ("[10] Port Scan ", "Scan open ports on a host"),
    ("[11] PCAP      ", "Analyze a PCAP file (extract + scan all IPs/domains)"),
    ("[12] Exit      ", "Quit 404 Scanner"),
]

def show_menu():
    print()
    divider()
    print(f"{Color.BOLD}  SELECT SCAN TYPE{Color.RESET}\n")
    for icon, desc in MENU_OPTIONS:
        print(f"  {Color.CYAN}{icon}{Color.RESET}  {desc}")
    divider()

def get_choice() -> int:
    while True:
        try:
            choice = int(input(f"\n{Color.BOLD}Enter choice [1-{len(MENU_OPTIONS)}]: {Color.RESET}"))
            if 1 <= choice <= len(MENU_OPTIONS):
                return choice
            log_warn(f"Enter a number between 1 and {len(MENU_OPTIONS)}.")
        except ValueError:
            log_warn("Invalid input. Please enter a number.")

def run_scan(choice: int):
    print()
    result = None
    try:
        if   choice == 1:
            target = input(f"{Color.CYAN}Enter URL          : {Color.RESET}").strip()
            result = scan_url(target)
        elif choice == 2:
            target = input(f"{Color.CYAN}Enter IP Address   : {Color.RESET}").strip()
            result = scan_ip(target)
        elif choice == 3:
            target = input(f"{Color.CYAN}Enter Domain       : {Color.RESET}").strip()
            result = scan_domain(target)
        elif choice == 4:
            target = input(f"{Color.CYAN}Enter Hash (MD5/SHA-256): {Color.RESET}").strip()
            result = scan_hash(target)
        elif choice == 5:
            target = input(f"{Color.CYAN}Enter File Path    : {Color.RESET}").strip()
            result = scan_file(target)
        elif choice == 6:
            target = input(f"{Color.CYAN}Enter PDF Path     : {Color.RESET}").strip()
            result = scan_pdf(target)
        elif choice == 7:
            target = input(f"{Color.CYAN}Enter Image Path   : {Color.RESET}").strip()
            result = scan_image(target)
        elif choice == 8:
            target = input(f"{Color.CYAN}Enter Word Path    : {Color.RESET}").strip()
            result = scan_word(target)
        elif choice == 9:
            target = input(f"{Color.CYAN}Enter TXT Path     : {Color.RESET}").strip()
            result = scan_txt(target)
        elif choice == 10:
            target = input(f"{Color.CYAN}Enter Host/IP      : {Color.RESET}").strip()
            result = scan_ports(target)
        elif choice == 11:
            target = input(f"{Color.CYAN}Enter PCAP Path    : {Color.RESET}").strip()
            result = scan_pcap(target)
        if result:
            print_report(result)
            ask_save_report(result)
    except FileNotFoundError as e:
        log_error(str(e))
    except ImportError as e:
        log_error(str(e))
    except requests.RequestException as e:
        log_error(f"Network error: {e}")
    except Exception as e:
        log_error(f"Unexpected error: {e}")
    input(f"\n{Color.CYAN}Press Enter to return to menu...{Color.RESET}")


# ──────────────────────────────────────────────
# Entry Point
# ──────────────────────────────────────────────

def main():
    banner()
    log_success("No API key needed — loading threat intelligence feeds...\n")
    _load_openphish()   # pre-fetch OpenPhish feed once at startup
    while True:
        show_menu()
        choice = get_choice()
        if choice == len(MENU_OPTIONS):
            print(f"\n{Color.CYAN}Goodbye! Team 404 Scanner signing off.Stay Secure!🛡️\n{Color.RESET}")
            break
        run_scan(choice)

if __name__ == "__main__":
    main()
